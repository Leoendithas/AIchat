import streamlit as st
import sqlite3
from openai import OpenAI
import json
import os
from datetime import datetime
from docx import Document
import io
import streamlit.components.v1 as components  # For rendering custom HTML
from streamlit_autorefresh import st_autorefresh

st.set_page_config(initial_sidebar_state="collapsed")

#==================================================================================
# Set your OpenAI API key.
#==================================================================================
client = OpenAI(api_key=st.secrets["api_keys"]["openai"])

#==================================================================================
# Auto-refresh every 3 seconds
#==================================================================================
st_autorefresh(interval=2000, key="chat_refresh")

#==================================================================================
# Database functions for shared persistent storage using SQLite.
#==================================================================================
def get_db_connection():
    conn = sqlite3.connect("chat.db", check_same_thread=False)
    conn.row_factory = sqlite3.Row
    return conn

def init_db():
    conn = get_db_connection()
    with conn:
        conn.execute('''
            CREATE TABLE IF NOT EXISTS messages (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                user TEXT NOT NULL,
                content TEXT NOT NULL,
                timestamp DATETIME DEFAULT CURRENT_TIMESTAMP
            )
        ''')
    conn.close()

def insert_message(user, content):
    conn = get_db_connection()
    with conn:
        conn.execute("INSERT INTO messages (user, content) VALUES (?, ?)", (user, content))
    conn.close()

def get_all_messages():
    conn = get_db_connection()
    messages = conn.execute("SELECT * FROM messages ORDER BY id").fetchall()
    conn.close()
    return messages

def get_user_message_count():
    conn = get_db_connection()
    count = conn.execute("SELECT COUNT(*) FROM messages WHERE user != 'GPT4o'").fetchone()[0]
    conn.close()
    return count

def get_last_message():
    conn = get_db_connection()
    msg = conn.execute("SELECT * FROM messages ORDER BY id DESC LIMIT 1").fetchone()
    conn.close()
    return msg

def clear_conversation():
    conn = get_db_connection()
    with conn:
        conn.execute("DELETE FROM messages")
    conn.close()

def get_active_members():
    """Extracts unique usernames from the messages, excluding 'GPT4o'."""
    messages = get_all_messages()
    active_users = {msg["user"] for msg in messages if msg["user"].lower() != "gpt4o"}
    return active_users


#==================================================================================
# Function to call GPT-4 API.
#==================================================================================
discussion_topic = "Should there be aircon in school classrooms? Why or why not?"

def get_gpt_response(conversation_history):
    openai_messages = []
    for msg in conversation_history:
        if msg["user"].lower() in ["gpt4o", "assistant"]:
            openai_messages.append({"role": "assistant", "content": msg["content"]})
        else:
            openai_messages.append({"role": "user", "content": msg["content"]})
    
    # Here we build a prompt (you can adjust as needed)
    active_members = str(len(get_active_members()))

    
    chat_history = "\n".join([f"{msg['user']}: {msg['content']}" for msg in conversation_history])
    
    prompt = f"""# Context #
You are an expert at facilitation for middle school students. 
Your role is to analyse the chat history and provide a suitable response to guide students to have a focused and productive discussion
based on the criteria provided below.

The discussion topic is {discussion_topic} and there are {active_members} active members in the Chat room.

# Completion Steps #
1. Analyse the chat history and identify key themes.
2. Based on the chat history, determine which criteria would be the most suitable response to guide the discussion.
3. Generate a two sentence response:
Sentence 1: Provide the key themes of student's responses thus far.
Sentence 2: Provide guidance by asking a question based on one of the 6 criteria below.

Sample response:
Students, I’ve noticed some interesting ideas coming up. For instance, some of you mentioned that uniforms might help create a more unified school look, while others pointed out concerns about expressing your personal style. 
Can you share some specific ways you think wearing uniforms could affect our daily school life? Please give examples of both positive effects—like feeling more focused or united—and any challenges you foresee. I'd love to hear from everyone!

Only Output Step 3 in your response.

<criteria>
6 Criteria and Weightage
1. No Response 
* Stay Quiet: Offer no response as the discussion is going well or students need more time to develop the discussion. Return with a ' ' only. 

2. Engagement Boosters 
* Highlight Inactive Students: Identify and encourage participation from those who haven't contributed yet.
* Encourage Equitable Participation: Ensure every student has an equal opportunity to share their ideas.

3. Guidance Enhancements 
* Elicit Further Explanation from Students: Ask students to elaborate on their thoughts for deeper insight.
* Prompt Reflection on Ideas: Encourage students to think critically about the underlying meaning of their ideas.
* Ask Probing Follow-Up Questions: Pose challenging questions that push students to analyze their viewpoints further.

4a. Discussion Steering  
* Politely remind students to stay on topic and guide them back to the main discussion.

4b. Organising and Summarising Opinions 
* Organize/Group Diverse Opinions: Summarize and categorize varying perspectives for clearer understanding.
* Periodically Summarize Opinions: Regularly recap key points to reinforce the discussion and ensure clarity.

5. Clarification & Error Correction 
* Identify & Clarify Misconceptions: Detect inaccuracies and correct misunderstandings within the discussion.
* Encourage Self-Correction: Motivate students to recognize and amend their own errors.
* Ask Students to Reevaluate Their Answers: Prompt students to reconsider and possibly refine their responses.

6. Emotional and Motivational Support 
* Provide Positive Reinforcement: Acknowledge and praise students for their valuable contributions.
* Offer Emotional Assurance/Validation: Provide supportive feedback to boost confidence and create a safe discussion environment.
</criteria>

Please follow these principles when replying:
1. Style: Write in a clear and informative way, using simple English. Address the students directly.
2. Tone: Maintain a positive motivational tone. Do not use negative sentences.
3. Audience: Students, age 12 to 15

Here is the chat history = {chat_history}

Remember to analyse the chat history thoroughly before deciding which is the most appropriate response to give.
"""
    response = client.chat.completions.create(
            model="gpt-4o-mini",  
            messages=[
                {"role": "system", "content": "You are a helpful assistant."},
                {"role": "user", "content": prompt}
            ],
            max_tokens=800,
            temperature=0.1
        )
    return response.choices[0].message.content

#==================================================================================
# Function to export conversation to a Word document using python-docx.
#==================================================================================
def export_to_word(messages):
    document = Document()
    document.add_heading("Chat Conversation", level=1)
    for msg in messages:
        document.add_paragraph(f"{msg['user']}: {msg['content']}")
    f = io.BytesIO()
    document.save(f)
    f.seek(0)
    return f

#==================================================================================
# Initialize the database.
#==================================================================================
init_db()

st.title("Shared Multi-User Chat Demo")

#==================================================================================
# Front-End Toggle: Enable/Disable API Call after every 5 messages
#==================================================================================
# This toggle appears in the sidebar.
enable_api_call = st.sidebar.checkbox("Enable AI", value=False)


#==================================================================================
# Chat Window (Scrollable with Auto-Scroll)
#==================================================================================
st.subheader(str(discussion_topic))

# Inject CSS for styling the chat window.
st.markdown(
    """
    <style>
    .chat-window {
        height: 400px;
        overflow-y: auto;
        border: 1px solid #ddd;
        padding: 10px;
    }
    </style>
    """, unsafe_allow_html=True
)

# Create the chat container.
chat_container = st.empty()
def render_chat():
    with chat_container.container():
        chat_html = '<div id="chat-window" class="chat-window">'
        messages = [dict(row) for row in get_all_messages()]
        for msg in messages:
            chat_html += f"<p><strong>{msg['user']}</strong>: {msg['content']}</p>"
        chat_html += "</div>"
        st.markdown(chat_html, unsafe_allow_html=True)
        
        # Use st.components.v1.html to execute the JavaScript for auto-scrolling.
        components.html(
            """
            <script>
              function scrollToBottom() {{
                var chatWindow = document.getElementById("chat-window");
                if(chatWindow) {{
                  chatWindow.scrollTop = chatWindow.scrollHeight;
                }}
              }}
              // Call the function immediately
              scrollToBottom();
            </script>
            """,
            height=0,  # no visible height needed
            scrolling=False,
        )
render_chat()

# Allow user to set their username; stored in session state.
if "username" not in st.session_state:
    st.session_state["username"] = ""

if st.session_state["username"] == "":
    st.session_state["username"] = st.text_input("Enter your username", key="username_input")

if st.session_state["username"]:
    # Use a form for the message input so that hitting Enter submits the message.
    with st.form(key="message_form", clear_on_submit=True):
        message_input = st.text_input("Enter a message", key="message_input")
        submit = st.form_submit_button("Send")
    
    if submit:
        user_message = message_input.strip()
        if user_message:
            insert_message(st.session_state["username"], user_message)
            # After insertion, check user message count and (if toggle enabled) trigger GPT response every 5 user messages.
            user_count = get_user_message_count()
            last_msg = get_last_message()
            if enable_api_call and (user_count % 10 == 0) and last_msg["user"] != "GPT4o":
                with st.spinner("Calling GPT-4..."):
                    conversation = [dict(row) for row in get_all_messages()]
                    gpt_reply = get_gpt_response(conversation)
                    insert_message("GPT4o", gpt_reply)
                st.success("GPT-4 has responded!")
        render_chat()
            
#==================================================================================
# Display Active Members Count
#==================================================================================
active_members = get_active_members()
st.markdown(f"**Active Members: {len(active_members)}**")
st.markdown("Active Users: " + ", ".join(active_members))

#==================================================================================
# Conversation Controls: Clear and Export
#==================================================================================
col1, col2 = st.columns(2)
with col1:
    if st.button("Clear Conversation"):
        clear_conversation()
        st.rerun()
with col2:
    if st.button("Export Conversation to Word"):
        messages = [dict(row) for row in get_all_messages()]
        word_file = export_to_word(messages)
        st.download_button(
            label="Download Word Document",
            data=word_file,
            file_name="chat_conversation.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
